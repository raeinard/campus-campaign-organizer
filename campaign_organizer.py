#!/usr/bin/env python3
"""
Campus Campaign Organizer – Local Python version
Usage: python campaign_organizer.py input_file.csv
       or python campaign_organizer.py input_file.xlsx
"""

import sys
import os
import re
import zipfile
import io
from datetime import datetime
import pandas as pd
import matplotlib.pyplot as plt
from rapidfuzz import process, fuzz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ------------------------------------------------------------
# 1. Standardise residence names
# ------------------------------------------------------------
CANONICAL_RES = [
    "Res 1", "Res 2", "Res 3",
    "Res 1a", "Res 1b", "Res 1c",
    "Res 2a", "Res 2b",
    "Res 4b",
    "Res 5a", "Res 5b",
    "Madeira Isles"
]

manual_map = {
    "res1": "Res 1", "res 1": "Res 1", "1": "Res 1",
    "2000 beds res 1": "Res 1", "2000beds res1": "Res 1", "2k beds res 1": "Res 1",
    "2000 beds residence 1": "Res 1",
    "res2": "Res 2", "res 2": "Res 2", "2": "Res 2",
    "2000 beds res 2": "Res 2", "2000beds res2": "Res 2", "2k beds res 2": "Res 2",
    "res3": "Res 3", "res 3": "Res 3", "3": "Res 3",
    "2000 beds res 3": "Res 3", "2000beds res3": "Res 3", "2k beds res 3": "Res 3",
    "res1a": "Res 1a", "1a": "Res 1a", "res 1a": "Res 1a",
    "res1b": "Res 1b", "1b": "Res 1b",
    "res1c": "Res 1c", "1c": "Res 1c",
    "res2a": "Res 2a", "2a": "Res 2a",
    "res2b": "Res 2b", "2b": "Res 2b",
    "res4b": "Res 4b", "4b": "Res 4b",
    "res5a": "Res 5a", "5a": "Res 5a",
    "res5b": "Res 5b", "5b": "Res 5b",
    "madeira": "Madeira Isles", "madeira isles": "Madeira Isles",
    "madeiraires": "Madeira Isles"
}

def standardise_res(text):
    if pd.isna(text):
        return "Unknown"
    original = str(text).strip()
    text_lower = original.lower()
    if text_lower in manual_map:
        return manual_map[text_lower]
    m = re.search(r'res(?:idence)?\s*(\d+)([a-c]?)', text_lower)
    if m:
        num = m.group(1)
        letter = m.group(2)
        return f"Res {num}{letter}" if letter else f"Res {num}"
    match, score, _ = process.extractOne(text_lower, CANONICAL_RES, scorer=fuzz.ratio)
    if score >= 85:
        return match
    return original.title()

# ------------------------------------------------------------
# 2. Parse room/block/unit
# ------------------------------------------------------------
def parse_room_block_unit(row):
    text = row["Room no. & Block"]
    residence = row["Res_standard"]
    if pd.isna(text):
        return (None, None, None, None)
    text = str(text).strip()
    is_2000_beds = residence in ["Res 1", "Res 2", "Res 3"]
    
    # Special pattern for non‑2000 beds: e.g., E101 -> Block E, Floor 1, Room 101
    if not is_2000_beds:
        m = re.match(r'^([A-Z])(\d)(\d{2,3})$', text, re.IGNORECASE)
        if m:
            block = m.group(1).upper()
            floor = int(m.group(2))
            room = m.group(3)
            return (floor, block, room, None)
    
    # General parsing
    block_pattern = r'Block\s*([A-Z0-9]+)|B\s*([A-Z0-9]+)'
    floor_pattern = r'(?:Floor|Flr|F)\s*(\d+)|(\d+)(?:st|nd|rd|th)?\s*floor'
    room_pattern = r'Room\s*([A-Z0-9]+)|R\s*(\d+)|#(\d+)'
    unit_pattern = r'Unit\s*([A-Z0-9]+)|U\s*([A-Z0-9]+)'
    
    block = floor = room = unit = None
    m = re.search(block_pattern, text, re.IGNORECASE)
    if m:
        block = m.group(1) or m.group(2)
    m = re.search(floor_pattern, text, re.IGNORECASE)
    if m:
        floor = m.group(1) or m.group(2)
        if floor:
            floor = int(floor)
    m = re.search(room_pattern, text, re.IGNORECASE)
    if m:
        room = m.group(1) or m.group(2) or m.group(3)
    m = re.search(unit_pattern, text, re.IGNORECASE)
    if m:
        unit = m.group(1) or m.group(2)
    if not room:
        digits = re.findall(r'\b(\d{3,4})\b', text)
        if digits:
            room = digits[0]
    if not block:
        letter = re.search(r'\b([A-Z])\b', text)
        if letter:
            block = letter.group(1)
    return (floor, block, room, unit)

# ------------------------------------------------------------
# 3. Generate Word documents per school
# ------------------------------------------------------------
def generate_school_documents(campaign_df, output_dir="outputs"):
    os.makedirs(output_dir, exist_ok=True)
    zip_path = os.path.join(output_dir, "campaign_plans_by_school.zip")
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        schools = campaign_df['School'].unique()
        for school in schools:
            school_df = campaign_df[campaign_df['School'] == school].copy()
            doc = Document()
            title = doc.add_heading(f'Campus Campaign – {school}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()
            # Summary
            doc.add_heading('Summary: Students per Residence', level=1)
            summary = school_df.groupby('Residence')['Student_count'].sum().reset_index()
            summary.columns = ['Residence', 'Students']
            table = doc.add_table(rows=1+len(summary), cols=2)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Residence'
            hdr[1].text = 'Students'
            for i, row in summary.iterrows():
                cells = table.rows[i+1].cells
                cells[0].text = str(row['Residence'])
                cells[1].text = str(row['Students'])
            doc.add_paragraph()
            # Detailed route
            doc.add_heading('Detailed Door‑to‑Door Route', level=1)
            current_res = None
            current_block = None
            for _, row in school_df.iterrows():
                residence = row['Residence']
                block = row['Block']
                floor = row['Floor']
                if residence != current_res:
                    doc.add_heading(f'🏠 Residence: {residence}', level=2)
                    current_res = residence
                    current_block = None
                block_label = f'Block: {block}' if block != 'Unknown' else 'Block: (not specified)'
                if block != current_block:
                    doc.add_heading(block_label, level=3)
                    current_block = block
                floor_text = f'Floor {floor}' if floor != 0 else 'Floor: (not given)'
                doc.add_paragraph(f'📌 {floor_text} – {row["Student_count"]} student(s)', style='List Bullet')
                # Prepare lists
                students = row['Students'].split(', ') if row['Students'] else []
                whatsapps = row['WhatsApp_numbers'].split(', ') if row['WhatsApp_numbers'] else []
                rooms = row['Room_numbers'].split(', ') if row['Room_numbers'] else []
                units = row['Units'].split(', ') if row['Units'] else []
                avail = row['Availability'].split(', ') if row['Availability'] else []
                max_len = max(len(students), len(whatsapps), len(rooms), len(units), len(avail))
                students += [''] * (max_len - len(students))
                whatsapps += [''] * (max_len - len(whatsapps))
                rooms += [''] * (max_len - len(rooms))
                units += [''] * (max_len - len(units))
                avail += [''] * (max_len - len(avail))
                # Table
                sub_table = doc.add_table(rows=1+max_len, cols=5)
                sub_table.style = 'Light List Accent 2'
                hdr = sub_table.rows[0].cells
                hdr[0].text = 'Student Name'
                hdr[1].text = 'WhatsApp'
                hdr[2].text = 'Room No.'
                hdr[3].text = 'Unit No.'
                hdr[4].text = 'Availability'
                for i in range(max_len):
                    cells = sub_table.rows[i+1].cells
                    cells[0].text = students[i]
                    cells[1].text = whatsapps[i]
                    cells[2].text = rooms[i]
                    cells[3].text = units[i]
                    cells[4].text = avail[i]
                doc.add_paragraph()
            # Save to in‑memory buffer and add to zip
            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
            safe_name = school.replace('/', '_').replace(' ', '_')
            zip_file.writestr(f"campaign_{safe_name}.docx", doc_buffer.getvalue())
            print(f"✅ Added campaign_{safe_name}.docx")
    print(f"\n✅ Zip file saved to: {zip_path}")
    return zip_path

# ------------------------------------------------------------
# Main
# ------------------------------------------------------------
def main():
    if len(sys.argv) < 2:
        print("Usage: python campaign_organizer.py <input_file.csv|xlsx>")
        sys.exit(1)
    input_file = sys.argv[1]
    if not os.path.exists(input_file):
        print(f"Error: File '{input_file}' not found.")
        sys.exit(1)
    
    print(f"Loading {input_file}...")
    if input_file.endswith('.csv'):
        df = pd.read_csv(input_file)
    else:
        df = pd.read_excel(input_file)
    
    # Required columns (adjust if your column names differ)
    required_cols = ['Full Names', 'WhatsApp Number', 'School', 'Res', 'Room no. & Block', 'When are you available']
    for col in required_cols:
        if col not in df.columns:
            print(f"Warning: Column '{col}' not found. Please check your input.")
            sys.exit(1)
    
    print("Standardising residence names...")
    df["Res_standard"] = df["Res"].apply(standardise_res)
    
    print("Parsing room/block/unit...")
    df[["floor", "block", "room_number", "unit"]] = df.apply(
        parse_room_block_unit, axis=1, result_type="expand"
    )
    df["floor"] = df["floor"].fillna(0).astype(int)
    df["block"] = df["block"].fillna("Unknown")
    df["room_number"] = df["room_number"].fillna("")
    df["unit"] = df["unit"].fillna("")
    
    print("Grouping data...")
    grouped = df.groupby(["School", "Res_standard", "block", "floor"])
    campaign_rows = []
    for (school, res, block, floor), group in grouped:
        campaign_rows.append({
            "School": school,
            "Residence": res,
            "Block": block,
            "Floor": floor,
            "Student_count": len(group),
            "Students": ", ".join(group["Full Names"]),
            "WhatsApp_numbers": ", ".join(group["WhatsApp Number"].astype(str)),
            "Room_numbers": ", ".join(group["room_number"].astype(str)),
            "Units": ", ".join(group["unit"].astype(str)),
            "Availability": ", ".join(group["When are you available"].astype(str))
        })
    campaign_df = pd.DataFrame(campaign_rows)
    campaign_df = campaign_df.sort_values(["School", "Residence", "Block", "Floor"])
    
    # Generate documents
    zip_path = generate_school_documents(campaign_df)
    
    # Also save cleaned CSV
    cleaned_csv = "cleaned_responses.csv"
    df.to_csv(cleaned_csv, index=False)
    print(f"✅ Cleaned data saved to: {cleaned_csv}")
    
    # Optional: plot residence counts
    try:
        res_counts = df["Res_standard"].value_counts()
        res_counts.plot(kind="bar", title="Students per Residence")
        plt.tight_layout()
        plt.savefig("residence_counts.png")
        print("✅ Residence chart saved: residence_counts.png")
    except:
        pass
    
    print("\n🎉 All done! Open the outputs folder to find your campaign zip.")

if __name__ == "__main__":
    main()
